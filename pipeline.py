"""
Shared pipeline for the AI-Assisted Resume Portfolio Generator.

Supports multi-file extraction (PDF, DOCX, DOC, TXT, MD, RTF, JSON, PNG, JPG, WEBP, etc.),
multimodal Gemini vision processing, customizable API keys, and theme-able HTML generation.
"""

import io
import json
import os
from pathlib import Path

from jinja2 import Environment, FileSystemLoader, select_autoescape
import docx
from pypdf import PdfReader
from pypdf.errors import PyPdfError

from google import genai
from google.genai import types
from google.genai import errors as genai_errors

# --------------------------------------------------------------------------
# Configuration
# --------------------------------------------------------------------------
BASE_DIR = Path(__file__).resolve().parent
RESUME_FILE = BASE_DIR / "resume.txt"
TEMPLATE_FILE = "template.html"
STYLE_FILE = BASE_DIR / "style.css"
OUTPUT_FILE = BASE_DIR / "portfolio.html"

MIN_RESUME_LENGTH = 10               # minimum characters accepted for raw text
DEFAULT_MODEL = "gemini-2.5-flash"   # default model

EMPTY_PORTFOLIO = {
    "name": "",
    "headline": "",
    "summary": "",
    "skills": [],
    "education": [],   # [{institution, qualification, duration}]
    "experience": [],  # [{role, organization, duration, description}]
    "projects": [],    # [{title, description, technologies:[]}]
    "achievements": [],
    "contact": {"email": "", "phone": "", "linkedin": "", "github": "", "links": []},
}


class PortfolioError(Exception):
    """Raised for any handled failure in the pipeline. Message is safe to show to a user."""


# --------------------------------------------------------------------------
# Step 1: Multi-Format File Extraction (PDF, Word, Text, Images)
# --------------------------------------------------------------------------
def extract_from_file_bytes(file_bytes: bytes, filename: str, mime_type: str = "") -> dict:
    """Extract text or create an image Part from arbitrary file bytes."""
    ext = Path(filename).suffix.lower()

    # 1. PDF Documents
    if ext == ".pdf" or mime_type == "application/pdf":
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            if reader.is_encrypted:
                raise PortfolioError(f"'{filename}' is password-protected. Please upload an unlocked PDF.")
            pages_text = [page.extract_text() or "" for page in reader.pages]
            text = "\n".join(pages_text).strip()
            if text:
                return {"type": "text", "filename": filename, "content": text}
            raise PortfolioError(
                f"No text extracted from '{filename}'. If this is a scanned document photo, "
                "please upload it as an image file (PNG/JPG/WEBP)."
            )
        except PyPdfError as exc:
            raise PortfolioError(f"Could not read PDF '{filename}': {exc}") from exc

    # 2. Word Documents (.docx, .doc)
    if ext in [".docx", ".doc"] or "word" in mime_type.lower():
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))
            text = "\n".join(full_text).strip()
            if not text:
                raise PortfolioError(f"Word document '{filename}' is empty.")
            return {"type": "text", "filename": filename, "content": text}
        except Exception as exc:
            raise PortfolioError(f"Could not extract content from Word document '{filename}': {exc}") from exc

    # 3. Image Files (.png, .jpg, .jpeg, .webp, .bmp, .tiff)
    if ext in [".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"] or mime_type.startswith("image/"):
        image_mime = mime_type if mime_type.startswith("image/") else f"image/{ext.replace('.', '')}"
        if image_mime == "image/jpg":
            image_mime = "image/jpeg"
        image_part = types.Part.from_bytes(data=file_bytes, mime_type=image_mime)
        return {"type": "image", "filename": filename, "part": image_part}

    # 4. Text / Markdown / Code / JSON (.txt, .md, .rtf, .json, etc.)
    try:
        text = file_bytes.decode("utf-8", errors="ignore").strip()
        if text:
            return {"type": "text", "filename": filename, "content": text}
        raise PortfolioError(f"File '{filename}' is empty.")
    except Exception as exc:
        raise PortfolioError(f"Could not read file '{filename}': {exc}") from exc


def clean_resume_text(text: str) -> str:
    """Strip unnecessary spaces and drop blank lines before sending text to Gemini."""
    lines = [line.strip() for line in text.splitlines()]
    non_empty_lines = [line for line in lines if line]
    return "\n".join(non_empty_lines)


def validate_resume_text(cleaned_text: str) -> None:
    if not cleaned_text.strip():
        raise PortfolioError("The resume content is empty. Add resume details and try again.")


def load_resume_text() -> str:
    """Read, clean, and validate resume.txt for CLI compatibility."""
    if not RESUME_FILE.exists():
        raise PortfolioError(
            f"'{RESUME_FILE.name}' was not found in {BASE_DIR}. "
            "Create a resume.txt file next to main.py and add resume content."
        )
    raw = RESUME_FILE.read_text(encoding="utf-8", errors="ignore")
    cleaned = clean_resume_text(raw)
    validate_resume_text(cleaned)
    return cleaned


def write_resume_text(text: str) -> None:
    """Overwrite resume.txt with new content."""
    RESUME_FILE.write_text(text, encoding="utf-8")


# --------------------------------------------------------------------------
# Step 2: Build Multimodal Extraction Prompt
# --------------------------------------------------------------------------
def build_prompt(resume_text: str) -> str:
    return f"""You extract structured portfolio data from the provided resume/document content.

STRICT RULES:
- Use only information explicitly present in the document text/images provided below.
- Do NOT invent or guess skills, experience, projects, achievements, companies, dates, or links.
- If a piece of information is missing, use an empty string "" or an empty list [], never a placeholder such as "N/A" or "Unknown".
- Keep the professional summary concise (2-3 sentences) and strictly factual based on the content.
- Return valid JSON only. No markdown, no code fences, no explanation, no extra text.

Return JSON matching exactly this structure and key names:
{{
  "name": "string - full name",
  "headline": "string - short professional identity (e.g. role or field of study)",
  "summary": "string - concise, factual professional summary",
  "skills": ["string", "..."],
  "education": [
    {{"institution": "string", "qualification": "string", "duration": "string"}}
  ],
  "experience": [
    {{"role": "string", "organization": "string", "duration": "string", "description": "string"}}
  ],
  "projects": [
    {{"title": "string", "description": "string", "technologies": ["string", "..."]}}
  ],
  "achievements": ["string", "..."],
  "contact": {{
    "email": "string", "phone": "string", "linkedin": "string", "github": "string",
    "links": ["string", "..."]
  }}
}}

DOCUMENT CONTENT:
\"\"\"
{resume_text}
\"\"\"
"""


def build_multimodal_contents(extracted_items: list, extra_text: str = "") -> list:
    """Build contents payload for Gemini containing image Parts and structured text prompt."""
    contents = []
    text_blocks = []

    if extra_text and extra_text.strip():
        text_blocks.append(f"--- DIRECT USER TEXT INPUT ---\n{extra_text.strip()}")

    for item in extracted_items:
        if item["type"] == "image":
            contents.append(item["part"])
            text_blocks.append(f"--- IMAGE DOCUMENT ({item['filename']}) --- Extract text, skills, experience, and contact details from this image document.")
        elif item["type"] == "text":
            text_blocks.append(f"--- DOCUMENT ({item['filename']}) ---\n{item['content']}")

    combined_text = "\n\n".join(text_blocks)
    if not combined_text.strip() and not any(i["type"] == "image" for i in extracted_items):
        raise PortfolioError("No text or document images were provided. Please upload a file or type content.")

    prompt_text = build_prompt(combined_text)
    contents.append(prompt_text)
    return contents


# --------------------------------------------------------------------------
# Step 3: Call Gemini API
# --------------------------------------------------------------------------
def call_gemini(contents: list | str, api_key: str = None, model_name: str = None) -> str:
    key = (api_key or os.getenv("GEMINI_API_KEY") or "").strip()
    if not key:
        raise PortfolioError(
            "GEMINI_API_KEY is missing. Please enter your Gemini API Key in the settings panel or set it in your .env file."
        )

    model = model_name or os.getenv("GEMINI_MODEL", DEFAULT_MODEL)

    try:
        client = genai.Client(api_key=key)
        response = client.models.generate_content(
            model=model,
            contents=contents,
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                temperature=0.2,
            ),
        )
    except genai_errors.ClientError as exc:
        raise PortfolioError(
            f"Gemini API request rejected ({exc}). Please check your API key and chosen model ({model})."
        ) from exc
    except genai_errors.ServerError as exc:
        raise PortfolioError(f"Gemini server error ({exc}). Please try again in a moment.") from exc
    except genai_errors.APIError as exc:
        raise PortfolioError(f"Gemini API call failed: {exc}") from exc
    except Exception as exc:
        raise PortfolioError(f"Could not reach Gemini API: {exc}") from exc

    text = (response.text or "").strip()
    if not text:
        raise PortfolioError("Gemini returned an empty response. Please try again.")
    return text


# --------------------------------------------------------------------------
# Step 4: Parse + Validate JSON safely
# --------------------------------------------------------------------------
def parse_portfolio_json(raw_text: str) -> dict:
    text = raw_text.strip()
    if text.startswith("```"):
        text = text.strip("`")
        if text.lower().startswith("json"):
            text = text[4:]
        text = text.strip()

    try:
        data = json.loads(text)
    except json.JSONDecodeError as exc:
        raise PortfolioError(f"Gemini did not return valid JSON ({exc}). Please try again.") from exc

    if not isinstance(data, dict):
        raise PortfolioError("Gemini output was not a valid JSON portfolio object.")

    return normalize_portfolio(data)


def normalize_portfolio(data: dict) -> dict:
    """Fill in any missing or wrongly-typed fields with safe empty values."""
    portfolio = {}

    for key, empty_value in EMPTY_PORTFOLIO.items():
        value = data.get(key, empty_value)
        if type(value) is not type(empty_value):
            value = empty_value
        portfolio[key] = value

    contact = portfolio["contact"]
    normalized_contact = {}
    for key, empty_value in EMPTY_PORTFOLIO["contact"].items():
        value = contact.get(key, empty_value) if isinstance(contact, dict) else empty_value
        if type(value) is not type(empty_value):
            value = empty_value
        normalized_contact[key] = value
    portfolio["contact"] = normalized_contact

    portfolio["skills"] = [s for s in portfolio["skills"] if isinstance(s, str) and s.strip()]
    portfolio["achievements"] = [a for a in portfolio["achievements"] if isinstance(a, str) and a.strip()]
    portfolio["education"] = [e for e in portfolio["education"] if isinstance(e, dict)]
    portfolio["experience"] = [e for e in portfolio["experience"] if isinstance(e, dict)]
    portfolio["projects"] = [p for p in portfolio["projects"] if isinstance(p, dict)]
    for project in portfolio["projects"]:
        techs = project.get("technologies", [])
        project["technologies"] = [t for t in techs if isinstance(t, str) and t.strip()] if isinstance(techs, list) else []

    return portfolio


# --------------------------------------------------------------------------
# Step 5: Render HTML
# --------------------------------------------------------------------------
def render_portfolio(portfolio: dict, theme: str = "dark") -> str:
    env = Environment(
        loader=FileSystemLoader(str(BASE_DIR)),
        autoescape=select_autoescape(["html"]),
    )
    template = env.get_template(TEMPLATE_FILE)
    inline_css = STYLE_FILE.read_text(encoding="utf-8")
    style_tag = f"<style>{inline_css}</style>"
    return template.render(style_tag=style_tag, theme=theme, **portfolio)


# --------------------------------------------------------------------------
# High-Level Orchestration (CLI & Default Web)
# --------------------------------------------------------------------------
def generate_portfolio_html(api_key: str = None, model: str = None, theme: str = "dark") -> str:
    """Run full pipeline against current resume.txt and return HTML."""
    resume_text = load_resume_text()
    prompt = build_prompt(resume_text)
    raw_json = call_gemini(prompt, api_key=api_key, model_name=model)
    portfolio = parse_portfolio_json(raw_json)
    html = render_portfolio(portfolio, theme=theme)
    OUTPUT_FILE.write_text(html, encoding="utf-8")
    return html
